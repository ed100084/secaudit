"""Document parsing helpers for uploaded audit framework source files."""
from dataclasses import dataclass, field
from pathlib import Path
import re

import PyPDF2
from docx import Document as DocxDocument
from openpyxl import load_workbook


@dataclass
class ParsedDocument:
    text: str
    markdown: str
    diagnostics: dict = field(default_factory=dict)


def _line_cleanup(text: str) -> str:
    lines = []
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = re.sub(r"[ \t]+", " ", line).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:]
    parts = [
        "| " + " | ".join(cell.strip() for cell in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in body:
        parts.append("| " + " | ".join(cell.strip() for cell in row) + " |")
    return "\n".join(parts)


def parse_pdf_document(file_path: str) -> ParsedDocument:
    text_parts = []
    markdown_parts = []
    page_count = 0
    pages_with_text = 0
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        page_count = len(reader.pages)
        for idx, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            clean = _line_cleanup(text)
            if clean:
                pages_with_text += 1
                text_parts.append(clean)
                markdown_parts.append(f"## Page {idx}\n\n{clean}")
            else:
                markdown_parts.append(f"## Page {idx}\n\n[No extractable text]")
    text = "\n\n".join(text_parts)
    diagnostics = {
        "parser": "PyPDF2",
        "page_count": page_count,
        "pages_with_text": pages_with_text,
        "char_count": len(text),
        "suspected_scanned": page_count > 0 and pages_with_text / page_count < 0.5,
    }
    return ParsedDocument(text=text, markdown="\n\n".join(markdown_parts), diagnostics=diagnostics)


def parse_docx_document(file_path: str) -> ParsedDocument:
    doc = DocxDocument(file_path)
    text_parts = []
    markdown_parts = []
    table_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)
            markdown_parts.append(text)
    for table in doc.tables:
        table_count += 1
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
                text_parts.append(" | ".join(cells))
        md = _markdown_table(rows)
        if md:
            markdown_parts.append(md)
    text = "\n".join(text_parts)
    return ParsedDocument(
        text=text,
        markdown="\n\n".join(markdown_parts),
        diagnostics={"parser": "python-docx", "table_count": table_count, "char_count": len(text)},
    )


def parse_xlsx_document(file_path: str) -> ParsedDocument:
    workbook = load_workbook(file_path, read_only=True, data_only=True)
    text_parts = []
    markdown_parts = []
    sheet_count = 0
    try:
        for sheet_name in workbook.sheetnames:
            sheet_count += 1
            worksheet = workbook[sheet_name]
            markdown_parts.append(f"## {sheet_name}")
            text_parts.append(f"=== {sheet_name} ===")
            rows = []
            for row in worksheet.iter_rows(values_only=True):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(cells):
                    rows.append(cells)
                    text_parts.append(" | ".join(cells))
            if rows:
                markdown_parts.append(_markdown_table(rows))
    finally:
        workbook.close()
    text = "\n".join(text_parts)
    return ParsedDocument(
        text=text,
        markdown="\n\n".join(markdown_parts),
        diagnostics={"parser": "openpyxl", "sheet_count": sheet_count, "char_count": len(text)},
    )


def parse_document(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8")
        clean = _line_cleanup(text)
        return ParsedDocument(
            text=clean,
            markdown=clean,
            diagnostics={"parser": "plain-text", "char_count": len(clean)},
        )
    parsers = {
        ".pdf": parse_pdf_document,
        ".docx": parse_docx_document,
        ".xlsx": parse_xlsx_document,
        ".xls": parse_xlsx_document,
    }
    parser = parsers.get(suffix)
    if not parser:
        raise ValueError(f"不支援的檔案格式: {suffix}")
    return parser(str(path))


def parse_file(file_path: str) -> str:
    return parse_document(file_path).text


def _split_table_row(line: str) -> list[str]:
    if "|" not in line:
        return []
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [cell for cell in cells if cell and cell != "---"]


def _header_map(cells: list[str]) -> dict:
    mapping = {}
    for idx, cell in enumerate(cells):
        label = cell.lower()
        if any(k in label for k in ("編號", "項次", "no", "序")):
            mapping["item_no"] = idx
        elif any(k in label for k in ("頻率", "週期", "時程", "期限")):
            mapping["frequency"] = idx
        elif any(k in label for k in ("佐證", "證據", "文件", "紀錄", "evidence")):
            mapping["evidence"] = idx
        elif any(k in label for k in ("依據", "法源", "條文", "reference")):
            mapping["reference"] = idx
        elif any(k in label for k in ("權責", "負責", "承辦", "單位", "owner")):
            mapping["owner"] = idx
        elif any(k in label for k in ("應辦", "辦理事項", "要求", "控制", "事項", "內容")):
            mapping["requirement"] = idx
    return mapping


def _cell(cells: list[str], idx: int | None) -> str:
    if idx is None or idx >= len(cells):
        return ""
    return cells[idx].strip()


def _is_requirement_text(text: str) -> bool:
    if len(text) < 10:
        return False
    keywords = ("應", "須", "需", "辦理", "建立", "訂定", "保存", "留存", "檢討", "定期", "紀錄", "佐證")
    return any(k in text for k in keywords)


def _make_control(requirement: str, source_text: str, item_no: str = "", reference: str = "",
                  owner: str = "", frequency: str = "", evidence: str = "", sort_order: int = 0) -> dict:
    notes = []
    if owner:
        notes.append(f"權責：{owner}")
    if frequency:
        notes.append(f"頻率/時程：{frequency}")
    if evidence:
        notes.append(f"佐證：{evidence}")
    full_requirement = requirement.strip()
    if notes:
        full_requirement = f"{full_requirement}\n" + "\n".join(notes)
    item = item_no or requirement[:36]
    return {
        "domain": "應辦事項",
        "item": item[:80],
        "level": "",
        "requirement": full_requirement[:3000],
        "reference": reference[:200],
        "source_text": source_text[:3000],
        "sort_order": sort_order,
    }


def extract_requirement_controls(markdown: str, limit: int = 120) -> list[dict]:
    controls = []
    seen = set()
    header = None
    sort_order = 0
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("| ---"):
            continue
        cells = _split_table_row(line)
        if cells:
            mapping = _header_map(cells)
            if "requirement" in mapping:
                header = mapping
                continue
            if header:
                requirement = _cell(cells, header.get("requirement"))
                if _is_requirement_text(requirement):
                    sort_order += 1
                    control = _make_control(
                        requirement=requirement,
                        source_text=" | ".join(cells),
                        item_no=_cell(cells, header.get("item_no")),
                        reference=_cell(cells, header.get("reference")),
                        owner=_cell(cells, header.get("owner")),
                        frequency=_cell(cells, header.get("frequency")),
                        evidence=_cell(cells, header.get("evidence")),
                        sort_order=sort_order,
                    )
                    signature = (control["item"], control["requirement"])
                    if signature not in seen:
                        seen.add(signature)
                        controls.append(control)
            elif any(_is_requirement_text(cell) for cell in cells):
                requirement = max(cells, key=len)
                sort_order += 1
                control = _make_control(requirement, " | ".join(cells), sort_order=sort_order)
                signature = (control["item"], control["requirement"])
                if signature not in seen:
                    seen.add(signature)
                    controls.append(control)
            if len(controls) >= limit:
                break
            continue

        match = re.match(r"^([0-9一二三四五六七八九十]+[.)、．]?|第[一二三四五六七八九十0-9]+[條項款點]?)[ \t]*(.+)$", line)
        item_no = ""
        requirement = line
        if match:
            item_no = match.group(1).strip()
            requirement = match.group(2).strip()
        if _is_requirement_text(requirement):
            sort_order += 1
            control = _make_control(requirement, line, item_no=item_no, sort_order=sort_order)
            signature = (control["item"], control["requirement"])
            if signature not in seen:
                seen.add(signature)
                controls.append(control)
        if len(controls) >= limit:
            break
    return controls


def build_compact_text(markdown: str, controls: list[dict], max_chars: int = 4000) -> str:
    if controls:
        lines = ["應辦事項摘要："]
        for idx, control in enumerate(controls[:40], 1):
            ref = f"（{control['reference']}）" if control.get("reference") else ""
            lines.append(f"{idx}. {control['item']}{ref}: {control['requirement'].replace(chr(10), ' ')}")
        compact = "\n".join(lines)
        if len(compact) >= max_chars:
            return compact[:max_chars]
        remaining = max_chars - len(compact) - 2
        return f"{compact}\n\n原文節錄：\n{markdown[:remaining]}"
    return markdown[:max_chars]
